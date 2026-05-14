import csv, os, time, shutil
from selenium import webdriver
from selenium.webdriver.common.by import By
from docx import Document
from docx.shared import Inches

base_dir = os.path.dirname(os.path.abspath(__file__))
logs_dir = os.path.join(base_dir, "..", "logs")
screenshots_dir = os.path.join(base_dir, "..", "screenshots")
reports_dir = os.path.join(base_dir, "..", "reports")

# Clean old evidence
for folder in [logs_dir, screenshots_dir, reports_dir]:
    if os.path.exists(folder):
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except Exception as e:
                print(f"Failed to delete {file_path}: {e}")
    os.makedirs(folder, exist_ok=True)

log_path = os.path.join(logs_dir, "browser_test_log.csv")
data_path = os.path.join(base_dir, "..", "data", "login_test_data.csv")
report_path = os.path.join(reports_dir, "Fail_Report.docx")

# Chrome setup
options = webdriver.ChromeOptions()
# options.add_argument("--headless")  # uncomment for headless mode
driver = webdriver.Chrome(options=options)
driver.implicitly_wait(5)
driver.get("file:///" + os.path.join(base_dir, "..", "demo_site.html"))

fail_cases_data, fail_cases_log = [], []

with open(data_path, newline="") as datafile, open(log_path, "w", newline="") as logfile:
    reader = csv.DictReader(datafile)
    writer = csv.writer(logfile)
    writer.writerow(["TestCaseID","Username","Password","Expected","Actual","Timestamp","Screenshot","Result"])
    case_num = 1
    for row in reader:
        username, password, expected = row["username"], row["password"], row["expected"]
        driver.find_element(By.ID, "username").clear()
        driver.find_element(By.ID, "password").clear()
        driver.find_element(By.ID, "username").send_keys(username)
        driver.find_element(By.ID, "password").send_keys(password)
        driver.find_element(By.TAG_NAME, "button").click()
        time.sleep(1)
        actual = driver.find_element(By.ID, "result").text
        timestamp = time.strftime("%Y%m%d_%H%M%S")
        screenshot_path = os.path.join(screenshots_dir, f"TC{case_num}_{username}_{timestamp}.png")
        driver.save_screenshot(screenshot_path)
        result = "Pass" if actual == expected else "Fail"
        writer.writerow([f"TC{case_num}", username, password, expected, actual, timestamp, screenshot_path, result])
        if result == "Fail":
            fail_cases_data.append([f"TC{case_num}", username, password, expected])
            fail_cases_log.append([f"TC{case_num}", username, password, expected, actual, timestamp, screenshot_path])
        case_num += 1

driver.quit()

doc = Document()
doc.add_heading("Fail Case Report", 0)
total_cases = case_num - 1
total_fail = len(fail_cases_log)
total_pass = total_cases - total_fail
doc.add_heading("Summary", level=1)
table = doc.add_table(rows=2, cols=3)
table.style = "Light List Accent 1"
hdr_cells = table.rows[0].cells
hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = "Total Cases","Pass","Fail"
row_cells = table.rows[1].cells
row_cells[0].text, row_cells[1].text, row_cells[2].text = str(total_cases), str(total_pass), str(total_fail)
doc.add_heading("Section 1: Fail Cases from login_test_data.csv", level=1)
for case in fail_cases_data:
    doc.add_paragraph(f"TestCaseID: {case[0]}, Username: {case[1]}, Password: {case[2]}, Expected: {case[3]}")
doc.add_heading("Section 2: Fail Cases from browser_test_log.csv", level=1)
for case in fail_cases_log:
    doc.add_paragraph(f"TestCaseID: {case[0]}, Username: {case[1]}, Password: {case[2]}, Expected: {case[3]}, Actual: {case[4]}, Timestamp: {case[5]}")
doc.add_heading("Section 3: Screenshots of Fail Cases", level=1)
for case in fail_cases_log:
    screenshot_path = case[6]
    if os.path.exists(screenshot_path):
        doc.add_paragraph(f"{case[0]} - {case[1]}")
        doc.add_picture(screenshot_path, width=Inches(4))
doc.save(report_path)
print(f"Word report created: {report_path}")
